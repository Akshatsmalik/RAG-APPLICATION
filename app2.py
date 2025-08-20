import os
import json
from dotenv import load_dotenv
from langchain_community.document_loaders import Docx2txtLoader, TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_groq import ChatGroq
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from docx import Document
from io import BytesIO
import re
from langchain import memory as memodule
from langchain.memory import ConversationBufferMemory
from langchain_ollama import OllamaLLM

load_dotenv()

DB_PATH = "faiss_index"

# ---------------- RAG Setup ---------------- #
def load_files(folder_path):
    docs = []
    print(f"Loading documents from {folder_path} ...")
    for file in os.listdir(folder_path):
        full_path = os.path.join(folder_path, file)
        if file.lower().endswith(".pdf"):
            loader = PyPDFLoader(full_path)
        elif file.lower().endswith(".txt"):
            loader = TextLoader(full_path, encoding="utf-8")
        elif file.lower().endswith(".docx"):
            loader = Docx2txtLoader(full_path)
        else:
            continue
        docs.extend(loader.load())
    print(f"Total documents loaded: {len(docs)}")
    return docs

def chunk_documents(docs):
    print("1")
    splitter = RecursiveCharacterTextSplitter(chunk_size=3000, chunk_overlap=300)
    chunks = splitter.split_documents(docs)
    print(f"Total chunks created: {len(chunks)}")
    return chunks

def save_faiss(vectordb, path):
    print("fiass")
    vectordb.save_local(path)
    print(f"FAISS index saved at {path}")

def load_faiss(path, embeddings):
    print(f"Loading FAISS index from {path} ...")
    print("load fiass")
    return FAISS.load_local(path, embeddings, allow_dangerous_deserialization=True)

def create_or_load_vectordb(chunks):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    if os.path.exists(DB_PATH) and os.listdir(DB_PATH) and chunks is None:
        vectordb = load_faiss(DB_PATH, embeddings)
        print("cl")
    else:
        vectordb = FAISS.from_documents(chunks, embeddings)
        save_faiss(vectordb, DB_PATH)
        print("fiass")
    return vectordb

# ---------------- LLM & QA Chain ---------------- #
from langchain_ollama import OllamaLLM

def chatbot_chain(vectordb):
    llm = OllamaLLM(
        model="llama3",  # your Ollama model name
        temperature=0,
    )

    retriever = vectordb.as_retriever(search_kwargs={"k": 3})

    prompt_template = PromptTemplate(
        input_variables=["context", "question"],
        template="""
You are "Corporate Agent", an ADGM legal compliance assistant.
Analyze the clause and return ONLY valid JSON with:
- compliance: "yes"/"no"/"uncertain"
- justification: detailed reason
- citations: list of exact ADGM law/regulation text from context
- red_flags: list of issues (invalid clauses, wrong jurisdiction, ambiguous language, non-compliance)
- suggested_text: short alternative wording
Context:
{context}
Question:
{question}
"""
    )

    chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True,
        chain_type_kwargs={"prompt": prompt_template},
    )
    return chain

# ---------------- Document Review ---------------- #
def extract_paragraphs(docx_bytes):
    print("chalu hogaya")
    doc = Document(BytesIO(docx_bytes))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def annotate_docx(original_bytes, annotations):
    doc = Document(BytesIO(original_bytes))
    for p in doc.paragraphs:
        for ann in annotations:
            if ann["text"] == p.text.strip():
                comment = f" [{ann['severity'].upper()}] {ann['note']}"
                p.add_run(comment).italic = True
    output = BytesIO()
    doc.save(output)
    return output.getvalue()

def review_document(file_bytes, chain):
    paragraphs = extract_paragraphs(file_bytes)
    annotations = []
    findings = []

    for para in paragraphs:
        # Run QA
        result = chain.invoke({"query": para})
        try:
            parsed = json.loads(result["result"])  # changed from result["answer"]
        except:
            parsed = {"compliance": "uncertain", "justification": "Parsing error", "red_flags": [], "citations": [], "suggested_text": ""}

        # Simple rule check for jurisdiction
        red_flags = parsed.get("red_flags", [])
        if re.search(r"\bUAE\b", para, re.I) and "ADGM" not in para:
            red_flags.append("Jurisdiction mismatch: mentions UAE but not ADGM")

        # Build annotations
        severity = "warning" if red_flags else "ok"
        note = "; ".join(red_flags) if red_flags else "Reviewed - no issues"
        annotations.append({"text": para, "note": note, "severity": severity})

        # Collect finding
        findings.append({
            "paragraph": para,
            "analysis": parsed,
            "red_flags": red_flags
        })

    # Annotated docx
    reviewed_bytes = annotate_docx(file_bytes, annotations)

    # JSON report
    report = {
        "process_detected": "company_incorporation",  # could be detected via prompt
        "document_count_uploaded": 1,
        "missing_mandatory_documents": [],
        "findings": findings
    }

    return reviewed_bytes, report

# ---------------- Main ---------------- #
if __name__ == "__main__":
    # Step 1: Load / build FAISS from ADGM reference folder
    folder_path = r"C:\Users\iaksh\Desktop\ml workspacec\RAG CHATBOT\downloads"
    if os.path.exists(DB_PATH) and os.listdir(DB_PATH):
        vectordb = create_or_load_vectordb(None)
    else:
        documents = load_files(folder_path)
        chunks = chunk_documents(documents)
        vectordb = create_or_load_vectordb(chunks)

    # Step 2: Create QA chain
    chain = chatbot_chain(vectordb)

    # Step 3: Load client docx to review
    print("Enter the path to the client document (docx) to review:")
    client_path = input(r"")
    with open(client_path, "rb") as f:
        file_bytes = f.read()

    reviewed_docx, json_report = review_document(file_bytes, chain)

    # Step 4: Save outputs
    with open("reviewed.docx", "wb") as f:
        f.write(reviewed_docx)
    with open("report.json", "w", encoding="utf-8") as f:
        json.dump(json_report, f, indent=2)

    print("Review complete. Outputs: reviewed.docx & report.json")
