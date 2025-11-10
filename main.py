import os
import json
import time
import hashlib
from contextlib import contextmanager
from io import BytesIO
import streamlit as st
from dotenv import load_dotenv
from docx import Document
import pypdf
import pandas as pd
import langchain
from langchain_community.document_loaders import Docx2txtLoader, TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_core.prompts import ChatPromptTemplate
from langchain.chains.retrieval import create_retrieval_chain
from langchain_ollama.llms import OllamaLLM
from langchain.schema import Document as LangchainDocument
import torch
from langchain_groq import ChatGroq

# from langchain_google_genai import ChatGoogleGenerativeAI
load_dotenv()

GROQ_API_KEY = "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"

FAISS_INDEX_DIR = "faiss_index"
RESPONSE_TIMEOUT = 30

@contextmanager
def timer(description):
    start_time = time.time()
    yield
    elapsed_time = time.time() - start_time
    st.info(f"⏱️ {description}: {elapsed_time:.2f} seconds")

@st.cache_resource
def get_embeddings():
    return HuggingFaceEmbeddings(
        model_name="sentence-transformers/all-MiniLM-L6-v2",
        model_kwargs={
            "device": "cpu",
        },
        encode_kwargs={"batch_size": 32}
    )

embeddings = get_embeddings()

@st.cache_data
def extract_text_from_pdf(pdf_bytes):
    try:
        pdf_reader = pypdf.PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

@st.cache_data
def extract_text_from_docx(docx_bytes):
    try:
        doc = Document(BytesIO(docx_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

@st.cache_data
def extract_text_from_csv(csv_bytes):
    try:
        df = pd.read_csv(BytesIO(csv_bytes))
        text = df.to_string(index=False)
        return text
    except Exception as e:
        st.error(f"Error reading CSV: {str(e)}")
        return ""

def process_uploaded_file(uploaded_file):
    file_bytes = uploaded_file.read()
    file_name = uploaded_file.name
    file_type = uploaded_file.type
    
    text_content = ""
    
    if file_type == "application/pdf" or file_name.lower().endswith('.pdf'):
        text_content = extract_text_from_pdf(file_bytes)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.lower().endswith('.docx'):
        text_content = extract_text_from_docx(file_bytes)
    elif file_type == "text/plain" or file_name.lower().endswith('.txt'):
        text_content = file_bytes.decode('utf-8', errors='ignore')
    elif file_type == "text/csv" or file_name.lower().endswith('.csv'):
        text_content = extract_text_from_csv(file_bytes)
    else:
        st.warning(f"Unsupported file type: {file_type}")
        return None
    
    if text_content:
        return LangchainDocument(
            page_content=text_content,
            metadata={"source": file_name, "file_type": file_type}
        )
    return None

def optimized_chunking(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=800,
        chunk_overlap=100,
        separators=["\n\n", "\n", ". ", "! ", "? ", " "],
        length_function=len
    )
    return splitter.split_documents(documents)

def create_qa_chain(vector_store):
    # llm = OllamaLLM(
    #     model="llama3:instruct", 
    #     temperature=0.3,
    #     num_ctx=2048,
    #     num_predict=200,
    #     num_thread=8
    # )

    llm = ChatGroq(
        model="deepseek-r1-distill-llama-70b",
        api_key=GROQ_API_KEY,
        temperature=0,
        max_tokens=None,
        reasoning_format="parsed",
        timeout=None,
        max_retries=2,
        # other params...
    )
    
    retriever = vector_store.as_retriever(search_kwargs={"k": 3})
    
    prompt_template = """You are an intelligent assistant that can answer questions in two modes:

1. DOCUMENT MODE: When the context contains relevant information, use it to answer the question
2. GENERAL MODE: When the context is not relevant or empty, answer from your general knowledge

Context from documents: {context}

Question: {question}

Instructions:
- If the context contains relevant information for the question, use it as your primary source
- If the context is not relevant or doesn't contain useful information, answer from your general knowledge
- Always be helpful and provide accurate information
- If you're unsure, say so clearly

Answer:"""
    
    prompt = PromptTemplate(
        input_variables=["context", "question"],
        template=prompt_template
    )
    
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True,
        chain_type_kwargs={"prompt": prompt},
    )
    return qa_chain

def analyze_document(document, qa_chain):
    content_preview = document.page_content[:2000] if len(document.page_content) > 2000 else document.page_content
    
    analysis_query = f"""Analyze this document and provide:
1. Document Type: What kind of document is this?
2. Main Topics: List the 3-5 main topics covered
3. Key Points: Highlight the most important information
4. Summary: Brief overview of the document content

Document content: {content_preview}"""
    
    result = qa_chain.invoke({"query": analysis_query})
    return result["result"]

st.set_page_config(page_title="Document Chat Assistant", page_icon="📄", layout="wide")

st.title("📄 Smart Document Chat Assistant")
st.write("Upload a document for analysis and chat with it, or ask general questions!")

with st.sidebar:
    st.header("📤 Upload Document")
    uploaded_file = st.file_uploader(
        "Choose a file", 
        type=["pdf", "docx", "txt", "csv"],
        help="Upload PDF, DOCX, TXT, or CSV files"
    )
    
    if uploaded_file:
        st.success(f"✅ File uploaded: {uploaded_file.name}")
        st.info(f"📊 File size: {len(uploaded_file.getvalue())} bytes")
    
    st.divider()
    st.header("ℹ️ How to Use")
    st.write("""
    1. **Upload** a document using the file uploader
    2. **Analyze** the document to get an overview
    3. **Chat** with the document or ask general questions
    
    **Examples:**
    - "What is this document about?"
    - "Summarize the key points"
    - "What is the capital of France?" (general)
    - "Explain quantum physics" (general)
    """)

col1, col2 = st.columns([1, 1])

with col1:
    st.header("📋 Document Analysis")
    
    if uploaded_file:
        if st.button("🔍 Analyze Document", type="primary"):
            with st.spinner("Processing document..."):
                with timer("Document Processing"):
                    document = process_uploaded_file(uploaded_file)
                    
                    if document:
                        chunks = optimized_chunking([document])
                        vector_store = FAISS.from_documents(chunks, embeddings)
                        
                        st.session_state.vector_store = vector_store
                        st.session_state.qa_chain = create_qa_chain(vector_store)
                        st.session_state.document = document
                        
                        analysis = analyze_document(document, st.session_state.qa_chain)
                        
                        st.subheader("📊 Document Analysis Results")
                        st.write(analysis)
                        
                        with st.expander("📄 Document Preview (First 500 characters)"):
                            st.text(document.page_content[:500] + "..." if len(document.page_content) > 500 else document.page_content)
                        
                        st.success("✅ Document processed! You can now chat with it.")
                    else:
                        st.error("❌ Failed to process the document.")
    else:
        st.info("👆 Please upload a document to begin analysis.")

with col2:
    st.header("💬 Chat Assistant")
    st.write("Analyse the document first to chat!")
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    chat_container = st.container()
    with chat_container:
        for message in st.session_state.messages:
            with st.chat_message(message["role"]):
                st.markdown(message["content"])
    
    if prompt := st.chat_input("Ask about your document or any general question..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        
        with st.chat_message("user"):
            st.markdown(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                try:
                    if hasattr(st.session_state, 'qa_chain') and st.session_state.qa_chain:
                        with timer("Response Generation"):
                            result = st.session_state.qa_chain.invoke({"query": prompt})
                            response = result["result"]
                            
                            sources = result.get("source_documents", [])
                            if sources and any(prompt.lower() in source.page_content.lower() for source in sources):
                                response += f"\n\n📚 *Source: {st.session_state.document.metadata.get('source', 'Uploaded document')}*"
                    else:
                        response = """I don't have access to any document right now. To answer questions about specific documents, please upload a file first using the sidebar.

However, I can help with general questions! Feel free to ask me about:
- General knowledge topics
- Explanations of concepts  
- How-to questions
- And much more!

What would you like to know?"""
                    
                    st.markdown(response)
                    
                except Exception as e:
                    response = f"I encountered an error: {str(e)}. Please try rephrasing your question or upload a document first."
                    st.error(response)
        
        st.session_state.messages.append({"role": "assistant", "content": response})

if st.session_state.messages:
    if st.button("🗑️ Clear Chat History"):
        st.session_state.messages = []
        st.rerun()

st.markdown("---")
st.markdown("🤖 **Smart Document Chat Assistant** | Ask questions about your documents or general topics!")
st.caption("💡 Tip: Upload a document first for document-specific questions, or ask general questions anytime!")











