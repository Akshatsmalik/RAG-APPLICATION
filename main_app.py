import os
import json
import re
from io import BytesIO
import tempfile
import shutil
import streamlit as st
from dotenv import load_dotenv
from docx import Document
import PyPDF2
import pandas as pd
from langchain_community.document_loaders import Docx2txtLoader, TextLoader, PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_community.vectorstores import FAISS
from langchain.prompts import PromptTemplate
from langchain.chains import RetrievalQA
from langchain_ollama import OllamaLLM
from langchain.schema import Document as LangchainDocument

load_dotenv()

FAISS_INDEX_DIR = "faiss_index"
DOCUMENTS_FOLDER = r"C:\Users\iaksh\Desktop\ml workspacec\RAG CHATBOT\downloads"

def load_documents_from_folder(folder_path):
    loaded_docs = []
    if not os.path.exists(folder_path):
        return loaded_docs
        
    for filename in os.listdir(folder_path):
        full_path = os.path.join(folder_path, filename)
        if filename.lower().endswith(".pdf"):
            loader = PyPDFLoader(full_path)
        elif filename.lower().endswith(".txt"):
            loader = TextLoader(full_path, encoding="utf-8")
        elif filename.lower().endswith(".docx"):
            loader = Docx2txtLoader(full_path)
        else:
            continue
        try:
            loaded_docs.extend(loader.load())
        except Exception as e:
            st.warning(f"Could not load {filename}: {str(e)}")
    return loaded_docs

def extract_text_from_pdf(pdf_bytes):
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(pdf_bytes))
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading PDF: {str(e)}")
        return ""

def extract_text_from_docx(docx_bytes):
    try:
        doc = Document(BytesIO(docx_bytes))
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    except Exception as e:
        st.error(f"Error reading DOCX: {str(e)}")
        return ""

def extract_text_from_csv(csv_bytes):
    try:
        df = pd.read_csv(BytesIO(csv_bytes))
        text = df.to_string(index=False)
        return text
    except Exception as e:
        st.error(f"Error reading CSV: {str(e)}")
        return ""

def process_uploaded_files(uploaded_files):
    documents = []
    
    for uploaded_file in uploaded_files:
        file_bytes = uploaded_file.read()
        file_name = uploaded_file.name
        file_type = uploaded_file.type
        
        text_content = ""
        
        if file_type == "application/pdf" or file_name.lower().endswith('.pdf'):
            text_content = extract_text_from_pdf(file_bytes)
        elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" or file_name.lower().endswith('.docx'):
            text_content = extract_text_from_docx(file_bytes)
        elif file_type == "text/plain" or file_name.lower().endswith('.txt'):
            text_content = file_bytes.decode('utf-8', errors='ignore')
        elif file_type == "text/csv" or file_name.lower().endswith('.csv'):
            text_content = extract_text_from_csv(file_bytes)
        else:
            st.warning(f"Unsupported file type for {file_name}: {file_type}")
            continue
        
        if text_content:
            doc = LangchainDocument(
                page_content=text_content,
                metadata={"source": file_name, "file_type": file_type}
            )
            documents.append(doc)
    
    return documents

def process_user_text_input(text_input, source_name="User Input"):
    if text_input.strip():
        doc = LangchainDocument(
            page_content=text_input,
            metadata={"source": source_name, "file_type": "text/plain"}
        )
        return [doc]
    return []

def split_documents_into_chunks(documents):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=3000, 
        chunk_overlap=300,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    chunks = splitter.split_documents(documents)
    return chunks

def save_faiss_index(vector_store, path):
    vector_store.save_local(path)

def load_faiss_index(path, embedding_model):
    return FAISS.load_local(path, embedding_model, allow_dangerous_deserialization=True)

def get_or_create_vector_store(doc_chunks, force_rebuild=False):
    embeddings = HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2")
    
    if not force_rebuild and os.path.exists(FAISS_INDEX_DIR) and os.listdir(FAISS_INDEX_DIR) and doc_chunks is None:
        try:
            vector_store = load_faiss_index(FAISS_INDEX_DIR, embeddings)
            return vector_store
        except Exception as e:
            st.warning(f"Could not load existing index: {str(e)}. Creating new one.")
    
    if doc_chunks is None:
        docs = load_documents_from_folder(DOCUMENTS_FOLDER)
        if not docs:
            st.warning("No documents found in the documents folder.")
            return None
        doc_chunks = split_documents_into_chunks(docs)
    
    if not doc_chunks:
        st.error("No document chunks to create vector store.")
        return None
    
    vector_store = FAISS.from_documents(doc_chunks, embeddings)
    save_faiss_index(vector_store, FAISS_INDEX_DIR)
    return vector_store

def build_qa_chain(vector_store, analysis_type="General Analysis"):
    llm = OllamaLLM(
        model="llama3",
        temperature=0
    )
    
    retriever = vector_store.as_retriever(search_kwargs={"k": 5})
    
    template_map = {
        "Compliance Review": """
You are 'Corporate Agent', an ADGM legal compliance assistant.
Analyze the clause and return ONLY valid JSON with:
- compliance: "yes"/"no"/"uncertain"
- justification: detailed reason
- citations: list of exact ADGM law/regulation text from context
- red_flags: list of issues (invalid clauses, wrong jurisdiction, ambiguous language, non-compliance)
- suggested_text: short alternative wording
Context: {context}
Question: {question}
""",
        "Risk Assessment": """
You are a legal risk assessment specialist.
Analyze the content and provide a detailed risk assessment including:
- High risk factors
- Medium risk factors
- Low risk factors
- Mitigation strategies
- Overall risk score (1-10)
Context: {context}
Question: {question}
""",
        "Document Summary": """
You are a document analysis expert. Provide a comprehensive summary including:
- Main topics covered
- Key findings
- Important clauses or sections
- Recommendations
Context: {context}
Question: {question}
""",
        "General Analysis": """
You are an AI assistant that provides detailed analysis based on the context provided.
Context: {context}
Question: {question}
Answer:
"""
    }
    
    prompt = PromptTemplate(
        input_variables=["context", "question"],
        template=template_map.get(analysis_type, template_map["General Analysis"])
    )
    
    qa_chain = RetrievalQA.from_chain_type(
        llm=llm,
        chain_type="stuff",
        retriever=retriever,
        return_source_documents=True,
        chain_type_kwargs={"prompt": prompt},
    )
    return qa_chain

def extract_paragraphs_from_docx(docx_bytes):
    doc = Document(BytesIO(docx_bytes))
    return [p.text.strip() for p in doc.paragraphs if p.text.strip()]

def annotate_document(doc_bytes, annotations):
    doc = Document(BytesIO(doc_bytes))
    for para in doc.paragraphs:
        for ann in annotations:
            if ann["text"] == para.text.strip():
                comment = f" [{ann['severity'].upper()}] {ann['note']}"
                para.add_run(comment).italic = True
    output_stream = BytesIO()
    doc.save(output_stream)
    return output_stream.getvalue()

def review_docx_document(doc_bytes, qa_chain):
    paragraphs = extract_paragraphs_from_docx(doc_bytes)
    annotations = []
    findings = []
    
    for para in paragraphs:
        result = qa_chain.invoke({"query": para})
        try:
            parsed_json = json.loads(result["result"])
        except Exception:
            parsed_json = {
                "compliance": "uncertain",
                "justification": "Parsing error",
                "red_flags": [],
                "citations": [],
                "suggested_text": ""
            }
        
        red_flags = parsed_json.get("red_flags", [])
        if re.search(r"\bUAE\b", para, re.I) and "ADGM" not in para:
            red_flags.append("Jurisdiction mismatch: mentions UAE but not ADGM")
        
        severity = "warning" if red_flags else "ok"
        note = "; ".join(red_flags) if red_flags else "Reviewed - no issues"
        annotations.append({"text": para, "note": note, "severity": severity})
        
        findings.append({
            "paragraph": para,
            "analysis": parsed_json,
            "red_flags": red_flags,
        })
    
    annotated_docx = annotate_document(doc_bytes, annotations)
    
    report = {
        "process_detected": "document_review",
        "uploaded_documents": 1,
        "missing_documents": [],
        "findings": findings,
    }
    
    return annotated_docx, report

def analyze_text_compliance(text_content, qa_chain):
    paragraphs = [p.strip() for p in text_content.split('\n') if p.strip()]
    
    findings = []
    for para in paragraphs:
        if len(para) < 50:
            continue
            
        result = qa_chain.invoke({"query": f"Analyze this text for compliance: {para}"})
        
        try:
            parsed_json = json.loads(result["result"])
        except:
            parsed_json = {
                "compliance": "uncertain",
                "justification": result["result"],
                "red_flags": [],
                "citations": [],
                "suggested_text": ""
            }
        
        red_flags = parsed_json.get("red_flags", [])
        if re.search(r"\bUAE\b", para, re.I) and "ADGM" not in para:
            red_flags.append("Jurisdiction mismatch: mentions UAE but not ADGM")
        
        findings.append({
            "paragraph": para,
            "analysis": parsed_json,
            "red_flags": red_flags,
        })
    
    return findings

st.title("🏛️ ADGM Legal Compliance Document Analyzer")
st.write("Upload documents (PDF, DOCX, TXT, CSV) or enter text directly for comprehensive legal analysis.")

with st.sidebar:
    st.header("⚙️ Configuration")
    
    if st.button("🔄 Rebuild Knowledge Base"):
        if os.path.exists(FAISS_INDEX_DIR):
            shutil.rmtree(FAISS_INDEX_DIR)
        st.success("Knowledge base will be rebuilt on next query.")
        st.rerun()
    
    st.subheader("System Status")
    if os.path.exists(FAISS_INDEX_DIR) and os.listdir(FAISS_INDEX_DIR):
        st.success("✅ Knowledge base loaded")
    else:
        st.info("📚 Building knowledge base...")
    
    st.subheader("Analysis Settings")
    default_analysis_type = st.selectbox(
        "Default Analysis Type",
        ["General Analysis", "Compliance Review", "Risk Assessment", "Document Summary"]
    )
    
    st.subheader("Vector Store Settings")
    chunk_size = st.slider("Chunk Size", 1000, 5000, 3000)
    chunk_overlap = st.slider("Chunk Overlap", 100, 500, 300)

@st.cache_resource(show_spinner=False)
def initialize_system():
    vector_store = get_or_create_vector_store(None)
    if vector_store is None:
        return None, None
    qa_chain = build_qa_chain(vector_store, default_analysis_type)
    return vector_store, qa_chain

with st.spinner("Initializing system..."):
    vector_store, qa_chain = initialize_system()

if vector_store is None or qa_chain is None:
    st.error("❌ Could not initialize the system. Please check your documents folder and try again.")
    st.stop()

tab1, tab2, tab3, tab4 = st.tabs(["📄 File Upload", "✍️ Text Input", "💬 Chat", "📊 Batch Analysis"])

with tab1:
    st.header("Upload Documents")
    uploaded_files = st.file_uploader(
        "Choose files", 
        type=["pdf", "docx", "txt", "csv"],
        accept_multiple_files=True
    )
    
    analysis_type = st.selectbox(
        "Analysis Type",
        ["Compliance Review", "General Analysis", "Document Summary", "Risk Assessment"]
    )
    
    if uploaded_files:
        with st.spinner("Processing uploaded files..."):
            documents = process_uploaded_files(uploaded_files)
            
            if documents:
                chunks = split_documents_into_chunks(documents)
                temp_vector_store = FAISS.from_documents(chunks, HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"))
                temp_qa_chain = build_qa_chain(temp_vector_store, analysis_type)
                
                st.success(f"✅ Processed {len(uploaded_files)} file(s) successfully!")
                
                for i, (uploaded_file, doc) in enumerate(zip(uploaded_files, documents)):
                    with st.expander(f"📋 Analysis: {uploaded_file.name}"):
                        if analysis_type == "Compliance Review" and uploaded_file.name.lower().endswith('.docx'):
                            file_bytes = uploaded_file.getvalue()
                            annotated_docx, report = review_docx_document(file_bytes, temp_qa_chain)
                            
                            st.download_button(
                                label="📥 Download Annotated Document",
                                data=annotated_docx,
                                file_name=f"reviewed_{uploaded_file.name}",
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            )
                            
                            st.json(report)
                        elif analysis_type == "Compliance Review":
                            findings = analyze_text_compliance(doc.page_content, temp_qa_chain)
                            st.json({"findings": findings[:3]})
                        elif analysis_type == "Document Summary":
                            result = temp_qa_chain.invoke({"query": f"Provide a summary of this document: {doc.page_content[:2000]}"})
                            st.write(result["result"])
                        else:
                            result = temp_qa_chain.invoke({"query": f"Analyze this document: {doc.page_content[:2000]}"})
                            st.write(result["result"])
                
                if st.button("📥 Download Analysis Report"):
                    report_data = {
                        "analysis_type": analysis_type,
                        "files_processed": [f.name for f in uploaded_files],
                        "timestamp": str(pd.Timestamp.now()),
                        "total_documents": len(documents)
                    }
                    
                    st.download_button(
                        label="Download Report (JSON)",
                        data=json.dumps(report_data, indent=2),
                        file_name=f"analysis_report_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.json",
                        mime="application/json"
                    )

with tab2:
    st.header("Direct Text Input")
    
    text_input = st.text_area(
        "Enter your text for analysis:",
        height=200,
        placeholder="Paste your legal text, contract clause, or any content you want to analyze..."
    )
    
    if text_input:
        documents = process_user_text_input(text_input, "Direct Input")
        
        if documents:
            chunks = split_documents_into_chunks(documents)
            temp_vector_store = FAISS.from_documents(chunks, HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"))
            
            analysis_option = st.selectbox(
                "Choose analysis type:",
                ["Compliance Check", "Legal Review", "Risk Assessment", "General Analysis"]
            )
            
            temp_qa_chain = build_qa_chain(temp_vector_store, analysis_option)
            
            if st.button("🔍 Analyze Text"):
                with st.spinner("Analyzing text..."):
                    query_map = {
                        "Compliance Check": f"Check compliance for: {text_input}",
                        "Legal Review": f"Provide legal review of: {text_input}",
                        "Risk Assessment": f"Identify legal risks in: {text_input}",
                        "General Analysis": f"Analyze: {text_input}"
                    }
                    
                    result = temp_qa_chain.invoke({"query": query_map[analysis_option]})
                    
                    st.subheader("📋 Analysis Results")
                    
                    if analysis_option == "Compliance Check":
                        try:
                            parsed_json = json.loads(result["result"])
                            st.json(parsed_json)
                        except:
                            st.write(result["result"])
                    else:
                        st.write(result["result"])
                    
                    if result.get("source_documents"):
                        with st.expander("📚 Source References"):
                            for i, doc in enumerate(result["source_documents"]):
                                st.write(f"**Source {i+1}:** {doc.metadata.get('source', 'Unknown')}")
                                st.write(doc.page_content[:300] + "...")

with tab3:
    st.header("💬 Interactive Chat")
    
    if "messages" not in st.session_state:
        st.session_state.messages = []
    
    for message in st.session_state.messages:
        with st.chat_message(message["role"]):
            st.markdown(message["content"])
    
    if prompt := st.chat_input("Ask me anything about legal compliance..."):
        st.session_state.messages.append({"role": "user", "content": prompt})
        with st.chat_message("user"):
            st.markdown(prompt)
        
        with st.chat_message("assistant"):
            with st.spinner("Thinking..."):
                result = qa_chain.invoke({"query": prompt})
                response = result["result"]
                st.markdown(response)
                
                if result.get("source_documents"):
                    with st.expander("📚 Sources"):
                        for i, doc in enumerate(result["source_documents"]):
                            st.write(f"**Source {i+1}:** {doc.metadata.get('source', 'Unknown')}")
        
        st.session_state.messages.append({"role": "assistant", "content": response})

with tab4:
    st.header("📊 Batch Analysis")
    
    batch_files = st.file_uploader(
        "Upload multiple files for batch processing",
        type=["pdf", "docx", "txt", "csv"],
        accept_multiple_files=True,
        key="batch_upload"
    )
    
    batch_analysis_type = st.selectbox(
        "Batch Analysis Type",
        ["Compliance Review", "Risk Assessment", "Document Summary", "General Analysis"],
        key="batch_analysis_type"
    )
    
    if batch_files and st.button("🚀 Start Batch Analysis"):
        progress_bar = st.progress(0)
        status_text = st.empty()
        
        batch_results = []
        
        for i, uploaded_file in enumerate(batch_files):
            status_text.text(f"Processing {uploaded_file.name}...")
            
            documents = process_uploaded_files([uploaded_file])
            
            if documents:
                chunks = split_documents_into_chunks(documents)
                temp_vector_store = FAISS.from_documents(chunks, HuggingFaceEmbeddings(model_name="all-MiniLM-L6-v2"))
                temp_qa_chain = build_qa_chain(temp_vector_store, batch_analysis_type)
                
                if batch_analysis_type == "Compliance Review":
                    result = temp_qa_chain.invoke({"query": f"Provide compliance analysis for: {documents[0].page_content[:1000]}"})
                    try:
                        analysis_result = json.loads(result["result"])
                    except:
                        analysis_result = {"analysis": result["result"]}
                else:
                    result = temp_qa_chain.invoke({"query": f"Analyze this document: {documents[0].page_content[:1000]}"})
                    analysis_result = {"analysis": result["result"]}
                
                batch_results.append({
                    "filename": uploaded_file.name,
                    "file_type": uploaded_file.type,
                    "analysis_type": batch_analysis_type,
                    "result": analysis_result
                })
            
            progress_bar.progress((i + 1) / len(batch_files))
        
        status_text.text("✅ Batch analysis completed!")
        
        st.subheader("📈 Batch Results Summary")
        
        for result in batch_results:
            with st.expander(f"📄 {result['filename']}"):
                st.write("**File Type:**", result['file_type'])
                st.write("**Analysis Type:**", result['analysis_type'])
                st.write("**Results:**")
                if isinstance(result['result'], dict):
                    st.json(result['result'])
                else:
                    st.write(result['result'])
        
        if st.button("📥 Download Batch Results"):
            batch_report = {
                "batch_analysis_date": str(pd.Timestamp.now()),
                "analysis_type": batch_analysis_type,
                "total_files": len(batch_files),
                "results": batch_results
            }
            
            st.download_button(
                label="Download Batch Report",
                data=json.dumps(batch_report, indent=2),
                file_name=f"batch_analysis_{pd.Timestamp.now().strftime('%Y%m%d_%H%M%S')}.json",
                mime="application/json"
            )

st.markdown("---")
st.markdown("🏛️ **ADGM Legal Compliance Assistant** | Built with Streamlit & LangChain")